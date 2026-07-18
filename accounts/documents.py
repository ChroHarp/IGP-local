from copy import deepcopy
from io import BytesIO
from pathlib import Path

from django.conf import settings
from django.core.exceptions import ObjectDoesNotExist
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from mailmerge import MailMerge

from .models import EducationTransitionRecord, SchoolSetting, StudentStaffAssignment


TEMPLATE_PATH = Path(settings.BASE_DIR) / "accounts" / "docx_templates" / "igp-template.docx"
MEETING_TEMPLATE_PATH = Path(settings.BASE_DIR) / "accounts" / "docx_templates" / "igp-meeting-template.docx"


class IGPDocumentError(ValueError):
    pass


def validate_igp_document(plan):
    missing = []
    if not TEMPLATE_PATH.is_file():
        raise IGPDocumentError("找不到 IGP DOCX 範本。")
    if not plan.student.full_name.strip():
        missing.append("學生姓名")
    if not plan.academic_year.strip():
        missing.append("學年度")
    if not plan.overall_goal.strip():
        missing.append("年度目標")

    semesters = list(
        plan.semester_plans.order_by("semester").prefetch_related(
            "course_plans__teacher",
            "course_plans__learning_performances__rating",
        )
    )
    if not semesters:
        missing.append("至少一份學期計畫")
    for semester in semesters:
        label = semester.get_semester_display()
        if not semester.goals.strip():
            missing.append(f"{label}目標")
        for course in semester.course_plans.all():
            if course.is_template or not course.is_active:
                continue
            if not course.goals.strip():
                missing.append(f"{label}「{course.course_name}」課程目標")

    if missing:
        raise IGPDocumentError("無法產出，請先補齊：" + "、".join(missing))
    return semesters


def _value(value):
    return str(value or "").strip()


def _profile_for(student):
    try:
        return student.initial_igp_profile
    except ObjectDoesNotExist:
        return None


def _family_values(member):
    if not member:
        return ("", "", "", "", "", "")
    return (
        member.relationship,
        member.full_name,
        member.education_major,
        member.specialty,
        member.phone,
        member.organization_or_school,
    )


def _merge_values(plan, fields):
    student = plan.student
    profile = _profile_for(student)
    guardian = student.guardians.order_by("-is_primary", "pk").first()
    members = list(student.family_members.all())
    father = next((member for member in members if member.relationship == "父親"), None)
    mother = next((member for member in members if member.relationship == "母親"), None)
    others = [member for member in members if member not in {father, mother}][:3]
    others += [None] * (3 - len(others))
    award = student.award_records.order_by("pk").first()

    values = {field: "" for field in fields}
    mapped = {
        "學生姓名": student.full_name,
        "班別": student.class_name,
        "性別": student.get_gender_display(),
        "住家電話": student.home_phone,
        "住址": student.address,
        "法定代理人": guardian.full_name if guardian else "",
        "法定代理人連絡電話_公司": guardian.phone_work if guardian else "",
        "法定代理人連絡電話__手機": guardian.phone_mobile if guardian else "",
        "家庭文化特質": profile.family_culture if profile else "",
        "主要照顧者": profile.primary_caregiver if profile else "",
        "主要協助學習者": profile.learning_supporter if profile else "",
        "家庭經濟狀況": profile.household_economy if profile else "",
        "照顧者管教態度": profile.caregiving_style if profile else "",
        "與家人互動情形": profile.family_interaction if profile else "",
        "第一階段數學性向測驗分數": profile.math_aptitude_score if profile else "",
        "第一階段自然性向測驗分數": profile.science_aptitude_score if profile else "",
        "第二階段數學實作評量_T分數": profile.math_practical_t_score if profile else "",
        "第二階段自然實作評量_T分數": profile.science_practical_t_score if profile else "",
        "科學興趣": profile.science_interests if profile else "",
        "人文與藝術興趣": profile.arts_interests if profile else "",
        "其他興趣": profile.other_interests if profile else "",
        "填寫者": profile.completed_by if profile else "",
        "認知優勢特質": plan.cognitive_strengths or (profile.cognitive_strengths if profile else ""),
        "情意優勢特質": plan.emotional_strengths or (profile.emotional_strengths if profile else ""),
        "學科優勢能力": plan.academic_strengths or (profile.academic_strengths if profile else ""),
        "認知弱勢特質": plan.cognitive_needs or (profile.cognitive_needs if profile else ""),
        "情意弱勢特質": plan.emotional_needs or (profile.emotional_needs if profile else ""),
        "學科弱勢能力": plan.academic_needs or (profile.academic_needs if profile else ""),
        "獲獎日期": award.award_date if award else "",
        "競賽活動名稱": award.activity_name if award else "",
        "主辦單位": award.organizer if award else "",
        "獎項": award.award if award else "",
        "得獎類型": award.award_type if award else "",
    }
    birth_field = next((field for field in fields if field.startswith("出生年月日")), None)
    if birth_field:
        mapped[birth_field] = student.date_of_birth.strftime("%Y/%m/%d") if student.date_of_birth else ""

    for prefix, member in (("父親", father), ("母親", mother)):
        _, name, education, specialty, phone, organization = _family_values(member)
        mapped.update({
            f"{prefix}姓名": name,
            f"{prefix}畢業科系": education,
            f"{prefix}專長": specialty,
            f"{prefix}連絡電話": phone,
            f"{prefix}服務機關": organization,
        })

    for index, member in enumerate(others, start=1):
        relationship, name, education, specialty, phone, organization = _family_values(member)
        relationship_field = "稱謂" if index == 1 else f"稱謂{index - 1}"
        education_field = {
            1: "家人1畢業科系_在學免填",
            2: "家人2畢業科系在學免填",
            3: "家人3畢業科系_在學免填",
        }[index]
        mapped.update({
            relationship_field: relationship,
            f"家人{index}姓名": name,
            education_field: education,
            f"家人{index}專長": specialty,
            f"家人{index}連絡電話": phone,
            f"家人{index}服務機關就讀學校": organization,
        })

    for field, value in mapped.items():
        if field in values:
            values[field] = _value(value)
    return values


def _replace_paragraph(paragraph, text, *, underlined_text=""):
    template_rpr = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            template_rpr = deepcopy(run._r.rPr)
            break
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)

    run = paragraph.add_run(_value(text))
    if template_rpr is not None:
        run._r.insert(0, deepcopy(template_rpr))
    if underlined_text:
        adjusted = paragraph.add_run("\n" + _value(underlined_text))
        if template_rpr is not None:
            adjusted._r.insert(0, deepcopy(template_rpr))
        adjusted.underline = True


def _set_cell(cell, text, *, underlined_text=""):
    paragraph = cell.paragraphs[0]
    _replace_paragraph(paragraph, text, underlined_text=underlined_text)
    for extra in cell.paragraphs[1:]:
        cell._tc.remove(extra._p)


def _prevent_cell_wrap(cell):
    properties = cell._tc.get_or_add_tcPr()
    if properties.find(qn("w:noWrap")) is None:
        properties.append(OxmlElement("w:noWrap"))

def _set_row(row, values):
    seen = set()
    for cell, value in zip(row.cells, values):
        if cell._tc in seen:
            continue
        seen.add(cell._tc)
        _set_cell(cell, value)


def _unique_cells(row):
    cells = []
    seen = set()
    for cell in row.cells:
        if cell._tc not in seen:
            seen.add(cell._tc)
            cells.append(cell)
    return cells


def _find_table(document, *required_text):
    normalized = tuple(item.replace(" ", "") for item in required_text)
    for table in document.tables:
        text = _body_text(table._tbl)
        if all(item in text for item in normalized):
            return table
    return None


def _replace_matching_paragraph(document, marker, text):
    paragraph = next((item for item in document.paragraphs if marker in item.text), None)
    if paragraph is not None:
        _replace_paragraph(paragraph, text)
    return paragraph


def _clear_matching_paragraphs(document, markers):
    for paragraph in document.paragraphs:
        if any(marker in paragraph.text for marker in markers):
            _replace_paragraph(paragraph, "")

def _clear_data_rows(table, start):
    for row in table.rows[start:]:
        _set_row(row, [""] * len(row.cells))


def _resize_data_rows(table, start, count):
    count = max(1, count)
    while len(table.rows) - start > count:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) - start < count:
        table._tbl.append(deepcopy(table.rows[-1]._tr))
    _clear_data_rows(table, start)


def _set_table_page_behavior(table, header_rows, data_start):
    for row in table.rows[:header_rows]:
        properties = row._tr.get_or_add_trPr()
        if properties.find(qn("w:tblHeader")) is None:
            properties.append(OxmlElement("w:tblHeader"))

    for row in table.rows[data_start:]:
        properties = row._tr.get_or_add_trPr()
        if properties.find(qn("w:cantSplit")) is None:
            properties.append(OxmlElement("w:cantSplit"))


def _body_text(element):
    return "".join(element.itertext()).replace(" ", "")


def _course_header_tables(document):
    return [
        table for table in document.tables
        if table.rows and table.cell(0, 0).text.strip() == "學習領域"
    ]


def _course_performance_tables(document):
    return [
        table for table in document.tables
        if table.rows and table.cell(0, 0).text.strip() == "項次"
    ]


def _ensure_course_slots(document, count):
    needed = max(1, count)
    headers = _course_header_tables(document)
    body = document._element.body
    marker = next(
        child for child in body.iterchildren()
        if "（二）課表" in _body_text(child)
    )

    if needed < len(headers):
        children = list(body.iterchildren())
        start = children.index(headers[needed]._tbl)
        end = children.index(marker)
        for child in children[start:end]:
            body.remove(child)
    elif needed > len(headers):
        children = list(body.iterchildren())
        start = children.index(headers[-1]._tbl)
        end = children.index(marker)
        block = children[start:end]
        for _ in range(needed - len(headers)):
            page_break = OxmlElement("w:p")
            run = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run.append(br)
            page_break.append(run)
            marker.addprevious(page_break)
            for child in block:
                marker.addprevious(deepcopy(child))


def _assignment_name(student, role):
    assignment = (
        student.staff_assignments.filter(role=role, is_active=True)
        .select_related("staff")
        .order_by("-start_date", "-pk")
        .first()
    )
    return assignment.staff.full_name if assignment else ""


def _grade_for_year(student, current_year, target_year):
    if student.grade is None:
        return ""
    try:
        grade = student.grade - (int(current_year) - int(target_year))
    except (TypeError, ValueError):
        grade = student.grade
    return grade if grade > 0 else ""


def _class_code(grade, class_number, fallback=""):
    if grade and class_number is not None:
        return f"{int(grade)}{int(class_number):02d}"
    fallback = _value(fallback)
    if fallback and fallback.isdigit() and grade and fallback.startswith(str(grade)):
        return fallback
    return fallback


def _populate_header_options(document, plan):
    categories = split_choices(plan.student.gifted_categories)
    selected = "、".join(categories) if categories else "未填"
    marker = next((p for p in document.paragraphs if "資優類別" in p.text), None)
    if marker is not None:
        _replace_paragraph(marker, f"資優類別：{selected}")
    _clear_matching_paragraphs(document, ("一般智能", "創造能力", "學術性向"))


def split_choices(value):
    if not value:
        return []
    return [item.strip() for item in str(value).replace("，", "、").replace(",", "、").split("、") if item.strip()]


def _populate_history(document, plan):
    student = plan.student
    homeroom = _assignment_name(student, StudentStaffAssignment.Role.HOMEROOM_TEACHER)
    case_manager = _assignment_name(student, StudentStaffAssignment.Role.CASE_MANAGER)
    semesters = list(
        plan.student.igp_plans.order_by("academic_year", "pk")
        .prefetch_related("semester_plans")
    )
    histories = [
        (annual, semester)
        for annual in semesters
        for semester in annual.semester_plans.all()
    ][-8:]
    table = document.tables[0]
    _clear_data_rows(table, 1)
    for row, (annual, semester) in zip(table.rows[1:], histories):
        grade = semester.grade or _grade_for_year(student, plan.academic_year, annual.academic_year)
        _set_row(row, [
            f"{annual.academic_year}-{semester.semester}",
            semester.school_name,
            _class_code(grade, semester.class_number, student.class_name),
            homeroom if annual.pk == plan.pk else "",
            case_manager if annual.pk == plan.pk else "",
            "",
            "",
        ])

def _remove_redundant_basic_page_break(document):
    paragraph = document.tables[1]._tbl.getnext()
    if paragraph is None or paragraph.tag != qn("w:p"):
        return
    for br in paragraph.findall(".//" + qn("w:br")):
        if br.get(qn("w:type")) == "page":
            br.getparent().remove(br)

def _populate_basic_details(document, plan):
    student = plan.student
    profile = _profile_for(student)
    basic = document.tables[1]
    if student.email:
        _set_cell(basic.cell(2, 7), student.email)

    family_results = (
        profile.family_culture if profile else "",
        profile.primary_caregiver if profile else "",
        profile.learning_supporter if profile else "",
        profile.household_economy if profile else "",
        profile.caregiving_style if profile else "",
        profile.family_interaction if profile else "",
    )
    for row_index, value in zip(range(13, 19), family_results):
        cells = _unique_cells(basic.rows[row_index])
        if len(cells) > 1:
            _set_cell(cells[1], value)

    transitions = {
        record.stage: record
        for record in student.education_transition_records.order_by("stage", "pk")
    }
    transition_table = document.tables[2]
    stages = (
        (EducationTransitionRecord.Stage.GRADES_3_4, "國小 3-4年級"),
        (EducationTransitionRecord.Stage.GRADES_5_6, "國小 5-6年級"),
        (EducationTransitionRecord.Stage.JUNIOR_HIGH, "國中"),
    )
    for row, (stage, label) in zip(transition_table.rows[2:], stages):
        record = transitions.get(stage)
        if record is None:
            _set_row(row, [label, "", "", "", "", ""])
            continue
        services = split_choices(record.service_types)
        if record.other_service:
            services.append(f"其他：{record.other_service}")
        _set_row(row, [
            label,
            record.school_name,
            record.class_name,
            record.homeroom_teacher,
            record.gifted_case_manager,
            "、".join(services),
        ])

    assessments = list(student.assessments.prefetch_related("subscales").all())
    table = document.tables[4]
    _resize_data_rows(table, 2, len(assessments))
    _set_table_page_behavior(table, 2, 2)
    for row, assessment in zip(table.rows[2:], assessments):
        subscales = [f"{subscale.name}：{subscale.score}" for subscale in assessment.subscales.all()]
        result = "\n".join(filter(None, [assessment.result, *subscales, assessment.notes]))
        _set_row(row, [
            assessment.name,
            _roc_date(assessment.assessed_on),
            result,
        ])

    interest_table = document.tables[5]
    interest_values = [
        profile.science_interests if profile else "",
        profile.arts_interests if profile else "",
        profile.other_interests if profile else "",
    ]
    for cell, value in zip(_unique_cells(interest_table.rows[2]), interest_values):
        _set_cell(cell, value)

    awards = list(student.award_records.order_by("pk"))
    table = document.tables[6]
    _resize_data_rows(table, 3, len(awards))
    _set_table_page_behavior(table, 3, 3)
    for number, (row, award) in enumerate(zip(table.rows[3:], awards), start=1):
        _set_row(row, [number, award.award_date, award.activity_name, award.organizer, award.award, award.award_type, ""])

def _grade_label(grade):
    return {3: "三", 4: "四", 5: "五", 6: "六", 7: "七", 8: "八", 9: "九"}.get(grade, _value(grade))


def _populate_annual_analysis(document, plan):
    profile = _profile_for(plan.student)
    table = document.tables[8]
    _clear_data_rows(table, 2)
    histories = list(plan.student.igp_plans.order_by("academic_year", "pk"))[-5:]
    today = timezone.localdate()
    for row, history in zip(table.rows[2:], histories):
        grade = _grade_for_year(plan.student, plan.academic_year, history.academic_year)
        _set_row(row, [
            _grade_label(grade),
            f"{today.year - 1911}/{today.month}/{today.day}" if history.pk == plan.pk else "",
            profile.completed_by if profile and history.pk == plan.pk else "",
            history.cognitive_strengths,
            history.emotional_strengths,
            history.academic_strengths,
            history.cognitive_needs,
            history.emotional_needs,
            history.academic_needs,
        ])


    analysis = "\n".join((
        "優勢：",
        f"學習領域（數理）：{plan.strength_math_science}",
        f"學習領域（語文）：{plan.strength_language}",
        f"劣勢：{plan.weakness_analysis}",
        f"情意方面：{plan.affective_analysis}",
        plan.qualitative_analysis,
    ))
    _set_cell(_unique_cells(table.rows[5])[0], analysis)
    _set_cell(_unique_cells(table.rows[6])[0], "")

def _populate_course_needs(document, semesters):
    table = document.tables[9]
    learning = "\n".join(
        f"{semester.get_semester_display()}：{semester.learning_domains}"
        for semester in semesters if semester.learning_domains
    )
    special = "\n".join(
        f"{semester.get_semester_display()}：{semester.special_needs_courses}"
        for semester in semesters if semester.special_needs_courses
    )
    _set_cell(table.cell(0, 1), learning)
    _set_cell(table.cell(1, 1), special)


def _rating_text(performance):
    try:
        return performance.rating.get_rating_display()
    except ObjectDoesNotExist:
        return ""


def _populate_course(header, performance_table, semester, course):
    area = course.learning_domains or course.special_needs_courses
    if not area:
        area = semester.learning_domains or semester.special_needs_courses
    _set_cell(header.cell(0, 1), area)
    _set_cell(header.cell(0, 3), course.course_name)
    _set_cell(header.cell(0, 5), course.teacher.full_name if course.teacher else "")
    _set_cell(
        header.cell(1, 1),
        f"{semester.igp_plan.academic_year}-{semester.semester}\n"
        f"年度目標：{semester.igp_plan.overall_goal}\n"
        f"學期目標：{semester.goals}\n"
        f"課程目標：{course.goals}",
    )
    _set_cell(header.cell(2, 2), course.cognitive_adjustments)
    _set_cell(header.cell(3, 2), course.affective_support)
    _set_cell(header.cell(4, 2), course.skill_training)

    performances = list(course.learning_performances.all())
    _resize_data_rows(performance_table, 2, len(performances))
    _set_table_page_behavior(performance_table, 2, 2)
    for row, performance in zip(performance_table.rows[2:], performances):
        _set_cell(row.cells[0], performance.sort_order)
        _set_cell(row.cells[1], performance.description, underlined_text=performance.adjustment)
        _set_cell(row.cells[2], performance.assessment_methods)
        _set_cell(row.cells[3], "")
        _set_cell(row.cells[4], _rating_text(performance))


def _populate_courses(document, semesters):
    courses = [
        (semester, course)
        for semester in semesters
        for course in semester.course_plans.all()
        if course.is_active and not course.is_template
    ]
    _ensure_course_slots(document, len(courses))
    headers = _course_header_tables(document)
    performance_tables = _course_performance_tables(document)

    if not courses:
        _populate_course(headers[0], performance_tables[0], semesters[0], _BlankCourse())
        return
    for header, performance_table, (semester, course) in zip(headers, performance_tables, courses):
        _populate_course(header, performance_table, semester, course)


def _roc_date(value):
    if not value:
        return ""
    return f"{value.year - 1911}/{value.month}/{value.day}"

def _meeting_values(meeting):
    return {
        "{{academic_year}}": meeting.igp_plan.academic_year,
        "{{semester}}": meeting.semester,
        "{{meeting_type}}": meeting.get_meeting_type_display(),
        "{{meeting_date}}": meeting.meeting_date.strftime("%Y/%m/%d") if meeting.meeting_date else "",
        "{{meeting_time}}": (meeting.meeting_time.strftime("%H:%M") if hasattr(meeting.meeting_time, "strftime") else _value(meeting.meeting_time)),
        "{{location}}": meeting.location,
        "{{recorder}}": meeting.recorder,
        "{{attendees}}": meeting.attendees,
        "{{minutes}}": meeting.minutes,
    }


def _replace_placeholders(document, values):
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in _unique_cells(row):
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        replacement = paragraph.text
        for marker, value in values.items():
            replacement = replacement.replace(marker, _value(value))
        if replacement != paragraph.text:
            _replace_paragraph(paragraph, replacement)


def _populate_meeting(document, plan):
    meeting = plan.meetings.order_by("-meeting_date", "-pk").first()
    if meeting is None:
        return
    _replace_matching_paragraph(
        document,
        "IGP期初/末會議",
        f"{plan.academic_year}學年度第{meeting.semester}學期 {meeting.get_meeting_type_display()}",
    )
    date_text = meeting.meeting_date.strftime("%Y/%m/%d") if meeting.meeting_date else ""
    time_text = meeting.meeting_time.strftime("%H:%M") if hasattr(meeting.meeting_time, "strftime") else _value(meeting.meeting_time)
    _replace_matching_paragraph(
        document,
        "會議日期:",
        f"會議日期：{date_text}  時間：{time_text}  地點：{meeting.location}  記錄者：{meeting.recorder}",
    )
    _replace_matching_paragraph(document, "與會人員:", f"與會人員（簽到）：{meeting.attendees}")
    minutes_table = _find_table(document, "課程方面", "情意方面", "獨立研究")
    if minutes_table is not None:
        _set_cell(_unique_cells(minutes_table.rows[0])[0], meeting.minutes)


def _populate_placement_reviews(document, plan):
    records = list(plan.student.placement_review_records.order_by("recorded_on", "pk")[:2])
    table = _find_table(document, "評估需求說明", "評估結果概要敘述", "記錄人員")
    if table is None:
        return
    _resize_data_rows(table, 1, len(records))
    _set_table_page_behavior(table, 1, 1)
    for row, record in zip(table.rows[1:], records):
        _set_row(row, [
            _roc_date(record.recorded_on),
            record.needs_description,
            record.result_summary,
            record.recorder,
        ])
        _prevent_cell_wrap(_unique_cells(row)[0])

def _populate_counseling(document, plan):
    records = list(
        plan.student.counseling_records.filter(academic_year=plan.academic_year)
        .select_related("author")
        .order_by("recorded_on", "pk")
    )
    table = _find_table(document, "參與人員", "內容概要敘述", "處遇方式")
    if table is None:
        return
    _resize_data_rows(table, 1, len(records))
    _set_table_page_behavior(table, 1, 1)
    for row, record in zip(table.rows[1:], records):
        author = record.author.get_full_name() or record.author.username
        _set_row(row, [
            _roc_date(record.recorded_on),
            record.participants,
            record.event,
            record.summary,
            record.intervention,
            author,
        ])
        _prevent_cell_wrap(_unique_cells(row)[0])


class _BlankCourse:
    learning_domains = ""
    special_needs_courses = ""
    course_name = ""
    teacher = None
    goals = ""
    cognitive_adjustments = ""
    affective_support = ""
    skill_training = ""

    class _Performances:
        @staticmethod
        def all():
            return []

    learning_performances = _Performances()


def build_igp_docx(plan):
    semesters = validate_igp_document(plan)
    merged = BytesIO()
    with MailMerge(TEMPLATE_PATH, auto_update_fields_on_open="no") as template:
        template.merge(**_merge_values(plan, template.get_merge_fields()))
        template.write(merged)
    merged.seek(0)

    document = Document(merged)
    _populate_header_options(document, plan)
    _populate_history(document, plan)
    _remove_redundant_basic_page_break(document)
    _populate_basic_details(document, plan)
    _populate_annual_analysis(document, plan)
    _populate_course_needs(document, semesters)
    _populate_courses(document, semesters)
    _populate_meeting(document, plan)
    _populate_placement_reviews(document, plan)
    _populate_counseling(document, plan)

    output = BytesIO()
    document.save(output)
    return output.getvalue()

def build_igp_meeting_docx(meeting):
    if not MEETING_TEMPLATE_PATH.is_file():
        raise IGPDocumentError("找不到 IGP 會議 DOCX 範本。")
    document = Document(MEETING_TEMPLATE_PATH)
    _replace_placeholders(document, _meeting_values(meeting))
    output = BytesIO()
    document.save(output)
    return output.getvalue()
